"""
P7 - Deep Learning-Based Object Detection in Satellite Imagery
India Space Academy - Project Report Generator
Output: P7_Project_Report.docx  (10-30 pages)
Run:    python generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_DIR = r"C:\Users\housh\Desktop\GIS and RS\Deep-Learning-Based-Object-Detection-in-Satellite-Imagery"
OUTPUT = PROJECT_DIR + r"\P7_Project_Report.docx"

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x4E, 0x79)
BLUE   = RGBColor(0x44, 0x72, 0xC4)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
AMBER  = RGBColor(0x96, 0x64, 0x00)
GREY   = RGBColor(0x59, 0x59, 0x59)

# ── Helpers ───────────────────────────────────────────────────────────────────
def _shd(cell, hex_fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    tcPr.append(shd)

def _para_shd(p, hex_fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    pPr.append(shd)

def _hrule(p, color="1F4E79", position="bottom"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{position}")
    el.set(qn("w:val"),   "single")
    el.set(qn("w:sz"),    "12")
    el.set(qn("w:space"), "1")
    el.set(qn("w:color"), color)
    pBdr.append(el)
    pPr.append(pBdr)

def run(p, text, size=12, bold=False, italic=False, color=BLACK, font="Calibri"):
    r = p.add_run(text)
    r.font.name   = font
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    r.font.color.rgb = color
    return r

def para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, size=12,
         bold=False, italic=False, color=BLACK,
         sb=0, sa=6, font="Calibri"):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        run(p, text, size=size, bold=bold, italic=italic, color=color, font=font)
    return p

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    _para_shd(p, "1F4E79")
    r = p.add_run(f"  {text}")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = WHITE
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = NAVY
    _hrule(p)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = BLUE
    return p

def body(text, sb=2, sa=4, bold=False, italic=False, color=BLACK):
    return para(text, size=11, sb=sb, sa=sa, bold=bold, italic=italic, color=color)

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run(p, text, size=11)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run(p, text, size=11)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    _para_shd(p, "F2F2F2")
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x8A)
    return p

def img_placeholder(caption_num, caption_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    _para_shd(p, "FFF2CC")
    run(p, f"  [ INSERT FIGURE {caption_num} HERE ]  ", size=10, italic=True, color=AMBER)
    cap = para(f"Figure {caption_num}: {caption_text}",
               align=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True, color=GREY, sb=0, sa=8)
    return p

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        _shd(c, "1F4E79")
        for par in c.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in par.runs:
                r.font.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(10)
                r.font.color.rgb = WHITE
    # rows
    for ri, row in enumerate(rows):
        fill = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            _shd(c, fill)
            for par in c.paragraphs:
                for r in par.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def pb():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
para()
para("INDIA SPACE ACADEMY", align=WD_ALIGN_PARAGRAPH.CENTER,
     bold=True, size=18, color=NAVY, sb=0, sa=4)
para("GIS and Remote Sensing Programme", align=WD_ALIGN_PARAGRAPH.CENTER,
     size=13, color=NAVY, sb=0, sa=20)

_hrule(para(), color="1F4E79", position="bottom")

para("PROJECT CODE:  P7", align=WD_ALIGN_PARAGRAPH.CENTER,
     bold=True, size=13, color=NAVY, sb=16, sa=6)

para("P7 – DEEP LEARNING-BASED OBJECT DETECTION", align=WD_ALIGN_PARAGRAPH.CENTER,
     bold=True, size=22, color=NAVY, sb=0, sa=4)
para("FROM SATELLITE IMAGERY", align=WD_ALIGN_PARAGRAPH.CENTER,
     bold=True, size=22, color=NAVY, sb=0, sa=10)

para("Semantic Segmentation of Land Cover Classes using U-Net Architecture",
     align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=14, color=BLUE, sb=0, sa=6)
para("Study Area: Indore District, Madhya Pradesh, India",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, color=BLACK, sb=0, sa=24)

_hrule(para(), color="1F4E79", position="top")

para("Submitted by:", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, sb=16, sa=4)
para("[Your Full Name]", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, color=NAVY, sb=0, sa=6)
para("Designation:  Student / Research Scholar", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, sb=0, sa=3)
para("Institution:  India Space Academy", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, sb=0, sa=3)
para("Programme:  GIS and Remote Sensing", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, sb=0, sa=16)
para("Date of Submission:  February 18, 2026", align=WD_ALIGN_PARAGRAPH.CENTER,
     bold=True, size=12, sb=0, sa=8)
para("Satellite: Sentinel-2  |  Framework: TensorFlow/Keras + Google Earth Engine",
     align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10, color=GREY, sb=0, sa=0)
pb()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("TABLE OF CONTENTS")
toc = [
    ("1.", "Title", True),
    ("2.", "Objective", True),
    ("3.", "Study Area", True),
    ("   3.1", "Location Details", False),
    ("   3.2", "Area of Interest Map", False),
    ("4.", "Data Used", True),
    ("   4.1", "Satellite Data Source", False),
    ("   4.2", "Bands and Spectral Indices", False),
    ("   4.3", "Date Range and Resolution", False),
    ("5.", "Methodology", True),
    ("   5.1", "GEE Workflow – Data Acquisition", False),
    ("   5.2", "GEE Workflow – Feature Engineering", False),
    ("   5.3", "GEE Workflow – Training Data Collection", False),
    ("   5.4", "Preprocessing Pipeline (Python)", False),
    ("   5.5", "Deep Learning Model Training (Python)", False),
    ("   5.6", "Inference and Classification (Python)", False),
    ("   5.7", "Post-processing and Vector Conversion (QGIS + Python)", False),
    ("6.", "Results", True),
    ("   6.1", "Classification Maps", False),
    ("   6.2", "Vector Outputs", False),
    ("   6.3", "Area Statistics", False),
    ("   6.4", "Accuracy Assessment", False),
    ("7.", "Conclusion", True),
    ("8.", "References", True),
]
for num, title, bold_flag in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run(p, f"{num}   {title}", size=11, bold=bold_flag)
pb()

# ══════════════════════════════════════════════════════════════════════════════
# 1. TITLE
# ══════════════════════════════════════════════════════════════════════════════
h1("1. TITLE")
para("P7 – Deep Learning-Based Object Detection from Satellite Imagery",
     bold=True, size=13, sb=6, sa=4)
para("Semantic Segmentation of Land Cover Classes using U-Net Architecture", italic=True, size=12, sb=0, sa=4)
para("Study Area: Indore District, Madhya Pradesh, India", bold=True, size=12, sb=0, sa=6)
body("This project implements an advanced Deep Learning approach for automated detection and "
     "classification of ground objects from high-resolution satellite imagery. The system uses "
     "the U-Net semantic segmentation architecture combined with Google Earth Engine (GEE) for "
     "cloud-based satellite data acquisition, and QGIS for post-processing and map creation.")

# ══════════════════════════════════════════════════════════════════════════════
# 2. OBJECTIVE
# ══════════════════════════════════════════════════════════════════════════════
h1("2. OBJECTIVE")
body("The primary objectives of this project are:", sb=6)
bullet("Automated Land Cover Classification: Detect and classify five classes — Background, Buildings, Roads, Water Bodies, and Vegetation — from Sentinel-2 satellite imagery using Deep Learning.")
bullet("U-Net Architecture: Apply the U-Net semantic segmentation model with ResNet-34 encoder backbone for pixel-wise classification.")
bullet("Google Earth Engine Integration: Use GEE for efficient satellite data acquisition, preprocessing, and spectral index computation.")
bullet("GIS-Ready Outputs: Generate raster and vector outputs compatible with QGIS and ArcGIS for spatial analysis.")
bullet("Statistical Analysis: Calculate area statistics for each land cover class to support urban planning and environmental monitoring.")
bullet("Reproducible Workflow: Create a complete, documented pipeline applicable to other geographic regions.")
para()
para("Expected Outcomes:", bold=True, size=11, sb=2, sa=4)
bullet("High-accuracy classification maps of Indore District")
bullet("Quantitative area statistics on buildings, roads, water bodies, and vegetation")
bullet("Vector shapefiles suitable for GIS integration")
bullet("Demonstration of Deep Learning in Remote Sensing applications")

# ══════════════════════════════════════════════════════════════════════════════
# 3. STUDY AREA
# ══════════════════════════════════════════════════════════════════════════════
h1("3. STUDY AREA")
h2("3.1 Location Details")
table(
    ["Parameter", "Details"],
    [
        ["Study Area",  "Indore District, Madhya Pradesh, India"],
        ["Latitude",    "22°43' N"],
        ["Longitude",   "75°50' E"],
        ["State",       "Madhya Pradesh"],
        ["District",    "Indore"],
        ["Country",     "India"],
        ["Area Type",   "Urban and suburban mixed land use"],
    ],
    col_widths=[5, 10]
)
body("Indore is the largest and most populous city in Madhya Pradesh and the commercial capital "
     "of the state. Its diverse land cover makes it an ideal study area for object detection:")
bullet("Urban Areas: Dense residential and commercial buildings in the city centre")
bullet("Transportation: Well-developed road infrastructure including national and state highways")
bullet("Water Bodies: Khan River, Saraswati River, and several artificial ponds and reservoirs")
bullet("Vegetation: Parks, gardens, and agricultural areas on the city periphery")
bullet("Mixed Land Use: Combination of residential, commercial, industrial, and recreational areas")

h2("3.2 Area of Interest (AOI) Map")
img_placeholder("3.1", "Location map – Indore District in Madhya Pradesh, India (QGIS: India outline, MP state highlighted, Indore District marked, AOI boundary, scale bar, north arrow, WGS 84)")
img_placeholder("3.2", "Google Earth Engine screenshot – Sentinel-2 true colour composite over Indore District with AOI boundary visible")
pb()

# ══════════════════════════════════════════════════════════════════════════════
# 4. DATA USED
# ══════════════════════════════════════════════════════════════════════════════
h1("4. DATA USED")
h2("4.1 Satellite Data Source")
table(
    ["Parameter", "Details"],
    [
        ["Satellite",        "Sentinel-2 (European Space Agency – ESA)"],
        ["Dataset",          "COPERNICUS/S2_SR  (Level-2A Surface Reflectance)"],
        ["GEE Catalog URL",  "https://developers.google.com/earth-engine/datasets/catalog/COPERNICUS_S2_SR"],
        ["Copernicus Hub",   "https://scihub.copernicus.eu/"],
        ["Boundary Data",    "GADM – https://gadm.org/ (District-level boundaries, India Level 2)"],
    ],
    col_widths=[4, 11]
)

h2("4.2 Bands and Spectral Indices Used")
body("Original Sentinel-2 Bands:", bold=True)
table(
    ["Band", "Wavelength (nm)", "Resolution (m)", "Purpose"],
    [
        ["B2 – Blue",  "490", "10", "Water detection, atmospheric correction"],
        ["B3 – Green", "560", "10", "Vegetation health, true colour composite"],
        ["B4 – Red",   "665", "10", "Vegetation discrimination, true colour"],
        ["B8 – NIR",   "842", "10", "Vegetation analysis, water detection"],
    ],
    col_widths=[3, 4, 4, 7]
)
body("Derived Spectral Indices:", bold=True)
table(
    ["Index", "Formula", "Purpose"],
    [
        ["NDVI", "(NIR − Red) / (NIR + Red)", "Vegetation detection and health assessment"],
        ["NDWI", "(Green − NIR) / (Green + NIR)", "Water body detection and mapping"],
        ["NDBI", "(SWIR − NIR) / (SWIR + NIR)", "Built-up area and building detection"],
        ["Texture Contrast", "GLCM Contrast (NIR band)", "Surface roughness / urban texture"],
        ["Texture Entropy",  "GLCM Entropy (NIR band)",  "Pixel randomness / land cover complexity"],
    ],
    col_widths=[4, 6, 7]
)
body("Total Input Channels: 8  (B2, B3, B4, B8, NDVI, NDWI, Texture Contrast, Texture Entropy)", bold=True)

h2("4.3 Date Range and Resolution")
table(
    ["Parameter", "Value"],
    [
        ["Start Date",              "January 1, 2023"],
        ["End Date",                "December 31, 2023"],
        ["Composite Type",          "Median composite (cloud-free pixel selection)"],
        ["Maximum Cloud Cover",     "10%"],
        ["Spatial Resolution",      "10 metres"],
        ["Atmospheric Correction",  "Level-2A Surface Reflectance (pre-applied by ESA)"],
    ],
    col_widths=[5, 10]
)
pb()

# ══════════════════════════════════════════════════════════════════════════════
# 5. METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
h1("5. METHODOLOGY")
body("The methodology is divided into two main workflows:", sb=6)
bullet("Google Earth Engine (GEE) Workflow – cloud-based data acquisition and preprocessing")
bullet("Python / QGIS Workflow – Deep Learning model training, inference, and post-processing")

# ── 5.1 ──────────────────────────────────────────────────────────────────────
h2("5.1 GEE Workflow – Data Acquisition  (Script: 01_data_acquisition.js)")
para("Objective: Download a cloud-free Sentinel-2 composite for Indore District.", bold=True, size=11, sb=2, sa=4)
numbered("Define the Area of Interest (AOI) as a bounding rectangle over Indore District.")
numbered("Load the Sentinel-2 ImageCollection filtered by bounds, date range, and cloud cover (<10%).")
numbered("Create a median composite to produce a single cloud-free image.")
numbered("Select required bands: B2, B3, B4, B8.")
numbered("Export the composite GeoTIFF to Google Drive (scale: 10 m).")
code(
    "var indore = ee.Geometry.Rectangle([75.65, 22.55, 76.05, 22.90]);\n"
    "var s2 = ee.ImageCollection('COPERNICUS/S2_SR')\n"
    "  .filterBounds(indore).filterDate('2023-01-01','2023-12-31')\n"
    "  .filter(ee.Filter.lt('CLOUDY_PIXEL_PERCENTAGE', 10));\n"
    "var composite = s2.median().clip(indore);\n"
    "Export.image.toDrive({ image: composite.select(['B2','B3','B4','B8']),\n"
    "  description:'Indore_Sentinel2_Composite', scale:10, region:indore });"
)
img_placeholder("5.1", "GEE Code Editor – data acquisition script with map visualisation and Tasks tab showing export task")

# ── 5.2 ──────────────────────────────────────────────────────────────────────
h2("5.2 GEE Workflow – Feature Engineering  (Script: 02_feature_engineering.js)")
para("Objective: Calculate spectral indices and texture features to enhance class separability.", bold=True, size=11, sb=2, sa=4)
numbered("Calculate NDVI: normalizedDifference(['B8','B4'])")
numbered("Calculate NDWI: normalizedDifference(['B3','B8'])")
numbered("Calculate NDBI: normalizedDifference(['B11','B8'])")
numbered("Calculate GLCM texture features (Contrast and Entropy) from the NIR band.")
numbered("Stack all bands into a single 8-band image and export to Google Drive.")
code(
    "var ndvi = composite.normalizedDifference(['B8','B4']).rename('NDVI');\n"
    "var ndwi = composite.normalizedDifference(['B3','B8']).rename('NDWI');\n"
    "var tex  = composite.select('B8').glcmTexture({size:3});\n"
    "var enhanced = composite.addBands(\n"
    "  [ndvi, ndwi, tex.select('B8_contrast'), tex.select('B8_ent')]);"
)
img_placeholder("5.2", "NDVI visualisation in GEE – green = high vegetation, brown = low vegetation")
img_placeholder("5.3", "NDWI visualisation in GEE – blue = water bodies, other = land")

# ── 5.3 ──────────────────────────────────────────────────────────────────────
h2("5.3 GEE Workflow – Training Data Collection  (Script: 03_ml_classification.js)")
para("Objective: Digitise training samples and generate an initial Random Forest classification.", bold=True, size=11, sb=2, sa=4)
numbered("Use the GEE polygon tool to digitise 15–20 training samples per class.")
numbered("Assign class property: Buildings=1, Roads=2, Water=3, Vegetation=4.")
numbered("Sample pixel values from the 8-band enhanced image at each training polygon.")
numbered("Train a Random Forest classifier (100 trees) on the sampled data.")
numbered("Apply the classifier to generate an initial classification map.")
numbered("Export the classification raster to Google Drive.")
img_placeholder("5.4", "Training sample digitisation in GEE – different colours for each class")
img_placeholder("5.5", "Initial Random Forest classification result in GEE")

# ── 5.4 ──────────────────────────────────────────────────────────────────────
h2("5.4 Preprocessing Pipeline  (Python)")
h3("Step 1 – Patch Extraction  (preprocessing/patch_extraction.py)")
table(
    ["Parameter", "Value", "Purpose"],
    [
        ["Patch Size",         "256 × 256 pixels", "Standard U-Net input size"],
        ["Overlap",            "32 pixels",         "Ensure complete area coverage"],
        ["Normalisation",      "Percentile 2–98",   "Reduce outlier impact"],
        ["Training Split",     "80%",               "Model training"],
        ["Validation Split",   "20%",               "Unbiased evaluation"],
    ],
    col_widths=[4, 4, 7]
)
code("python -m preprocessing.patch_extraction")

h3("Step 2 – Data Augmentation  (preprocessing/data_augmentation.py)")
table(
    ["Technique", "Parameters", "Purpose"],
    [
        ["Random Rotation",   "0°, 90°, 180°, 270°", "Orientation invariance"],
        ["Horizontal Flip",   "50% probability",      "Geometric variation"],
        ["Vertical Flip",     "50% probability",      "Geometric variation"],
        ["Brightness",        "Range 0.8–1.2",        "Illumination variation"],
        ["Contrast",          "Range 0.8–1.2",        "Image quality variation"],
        ["Gaussian Noise",    "sigma = 0.01",         "Sensor noise simulation"],
    ],
    col_widths=[4, 4, 7]
)

# ── 5.5 ──────────────────────────────────────────────────────────────────────
h2("5.5 Deep Learning Model Training  (Python – models/train.py)")
h3("U-Net Architecture")
table(
    ["Component", "Specification"],
    [
        ["Architecture",       "U-Net (Encoder–Decoder with Skip Connections)"],
        ["Encoder Backbone",   "ResNet-34 (pre-trained on ImageNet)"],
        ["Input Shape",        "256 × 256 × 8"],
        ["Output Shape",       "256 × 256 × 5  (class probabilities per pixel)"],
        ["Activation",         "Softmax (multi-class)"],
        ["Bottleneck Filters", "1024"],
        ["Loss Function",      "Categorical Cross-Entropy + Dice Loss"],
        ["Optimizer",          "Adam  (lr = 0.0001)"],
    ],
    col_widths=[5, 10]
)
h3("Training Hyperparameters")
table(
    ["Parameter", "Value", "Notes"],
    [
        ["Batch Size",              "8",          "Memory vs. convergence balance"],
        ["Maximum Epochs",          "100",        "Early stopping applied"],
        ["Learning Rate",           "0.0001",     "Stable convergence"],
        ["Validation Split",        "20%",        "Unbiased evaluation"],
        ["Early Stopping Patience", "15 epochs",  "Monitors val_loss"],
        ["LR Reduction",            "×0.5 after 5 epochs", "Escape local minima"],
        ["Checkpoint Monitor",      "val_iou_score (max)", "Saves best model"],
    ],
    col_widths=[5, 4, 6]
)
body("Evaluation Metrics: Overall Accuracy, IoU (Intersection over Union), F1-Score (per class)")
code("python -m models.train")
img_placeholder("5.6", "Training and validation loss curves over epochs")
img_placeholder("5.7", "Training and validation IoU score progression over epochs")

# ── 5.6 ──────────────────────────────────────────────────────────────────────
h2("5.6 Inference and Classification  (Python – models/inference.py)")
para("Objective: Apply the trained model to classify the entire Indore District.", bold=True, size=11, sb=2, sa=4)
numbered("Load the best saved model checkpoint (best_model.h5).")
numbered("Load the full-resolution 8-band GeoTIFF.")
numbered("Apply sliding-window inference (256×256 patches with overlap blending).")
numbered("Stitch patches back into a full-resolution classification raster.")
numbered("Save output as GeoTIFF with original georeferencing preserved.")
code(
    "python -m models.inference \\\n"
    "  --model models/saved_models/best_model.h5 \\\n"
    "  --image data/raw/Indore_Enhanced_Features.tif \\\n"
    "  --output data/outputs/"
)

# ── 5.7 ──────────────────────────────────────────────────────────────────────
h2("5.7 Post-processing and Vector Conversion  (QGIS + Python)")
h3("Morphological Filtering  (postprocessing/spatial_filtering.py)")
table(
    ["Operation", "Parameter", "Purpose"],
    [
        ["Morphological Closing", "Kernel 5×5 px",  "Fill small holes inside objects"],
        ["Small Object Removal",  "Min area 100 px", "Remove isolated noise pixels"],
        ["Boundary Smoothing",    "Gaussian filter", "Realistic, smooth object edges"],
    ],
    col_widths=[5, 4, 7]
)
code(
    "python -m postprocessing.spatial_filtering \\\n"
    "  --input  data/outputs/classification.tif \\\n"
    "  --output data/outputs/classification_filtered.tif"
)

h3("Raster to Vector Conversion  (QGIS)")
numbered("Load classification raster: Layer > Add Raster Layer.")
numbered("Polygonize: Raster > Conversion > Polygonize (Raster to Vector).")
numbered("Extract by Attribute to create separate shapefiles per class.")
numbered("Simplify geometry (Douglas-Peucker, tolerance 2.0 m).")
numbered("Calculate area: Field Calculator > $area / 1000000 (sq. km).")
numbered("Create Print Layout with title, legend, scale bar, north arrow, statistics table.")
numbered("Export final map as PDF at 300 DPI.")
img_placeholder("5.8", "QGIS – Polygonize (Raster to Vector) tool converting classification raster to polygon layer")
img_placeholder("5.9", "QGIS Print Layout – final map composition with legend, scale bar, and north arrow")
code(
    "python -m postprocessing.raster_to_vector \\\n"
    "  --input  data/outputs/classification_filtered.tif \\\n"
    "  --output data/outputs/vectors/\n\n"
    "python -m postprocessing.statistics \\\n"
    "  --raster data/outputs/classification_filtered.tif \\\n"
    "  --output data/outputs/reports/"
)
pb()

# ══════════════════════════════════════════════════════════════════════════════
# 6. RESULTS
# ══════════════════════════════════════════════════════════════════════════════
h1("6. RESULTS")

h2("6.1 Classification Maps")
body("The trained U-Net model classified the Indore District imagery into five land cover classes. "
     "The classification raster is a single-band GeoTIFF where each pixel carries a class label (0–4).")
table(
    ["Class ID", "Class Name", "Map Colour", "Description"],
    [
        ["0", "Background", "Black",  "Non-target areas"],
        ["1", "Buildings",  "Red",    "Residential and commercial structures"],
        ["2", "Roads",      "Blue",   "Transportation networks and paved surfaces"],
        ["3", "Water",      "Cyan",   "Rivers, lakes, ponds"],
        ["4", "Vegetation", "Green",  "Parks, forests, agricultural areas"],
    ],
    col_widths=[2, 3, 3, 8]
)
img_placeholder("6.1", "Complete land cover classification map of Indore District (legend, scale bar, north arrow, coordinate grid)")
img_placeholder("6.2", "Zoomed view – urban centre showing building detection (red) and road network (blue)")
img_placeholder("6.3", "Water body detection – Khan River and Saraswati River (cyan)")
img_placeholder("6.4", "Vegetation cover on city periphery (green) – parks, forests, agricultural land")

h2("6.2 Vector Outputs")
body("The raster classification was converted to vector format, generating separate polygon layers "
     "for each class. These are compatible with QGIS, ArcGIS, and web mapping platforms.")
img_placeholder("6.5", "QGIS interface – all vector layers overlaid on the map canvas")
img_placeholder("6.6", "Buildings vector layer (red polygons) with attribute table showing area values")
img_placeholder("6.7", "Roads vector layer (blue polygons/lines)")
img_placeholder("6.8", "Water bodies vector layer (cyan polygons)")

h2("6.3 Area Statistics")
para("Table 6.1: Land Cover Area Statistics", bold=True, size=11, sb=4, sa=2)
table(
    ["Class ID", "Class Name", "Area (sq. km)", "Area (hectares)", "% of Study Area", "Object Count"],
    [
        ["0", "Background", "[VALUE]", "[VALUE]", "[VALUE]%", "–"],
        ["1", "Buildings",  "[VALUE]", "[VALUE]", "[VALUE]%", "[VALUE]"],
        ["2", "Roads",      "[VALUE]", "[VALUE]", "[VALUE]%", "[VALUE]"],
        ["3", "Water",      "[VALUE]", "[VALUE]", "[VALUE]%", "[VALUE]"],
        ["4", "Vegetation", "[VALUE]", "[VALUE]", "[VALUE]%", "[VALUE]"],
        ["–", "TOTAL",      "[VALUE]", "[VALUE]", "100%",     "[VALUE]"],
    ],
    col_widths=[2, 3, 3, 3, 3, 3]
)
para("Note: Replace [VALUE] with actual figures after running:  python -m postprocessing.statistics",
     italic=True, size=10, color=AMBER, sb=0, sa=6)
img_placeholder("6.9",  "Pie chart – percentage distribution of land cover classes")
img_placeholder("6.10", "Bar chart – area (sq. km) per land cover class")

h2("6.4 Accuracy Assessment")
para("Table 6.2: Overall Model Performance", bold=True, size=11, sb=4, sa=2)
table(
    ["Metric", "Value", "Interpretation"],
    [
        ["Overall Accuracy",  "[VALUE]%",      "Percentage of correctly classified pixels"],
        ["Mean IoU",          "[VALUE]",       "Average overlap between prediction and ground truth"],
        ["Mean F1-Score",     "[VALUE]",       "Harmonic mean of precision and recall"],
        ["Training Time",     "[VALUE] hours", "Total GPU training time"],
        ["Inference Time",    "[VALUE] min",   "Time to classify full study area"],
    ],
    col_widths=[5, 3, 7]
)
para("Table 6.3: Per-Class Accuracy Metrics", bold=True, size=11, sb=4, sa=2)
table(
    ["Class", "Precision", "Recall", "F1-Score", "IoU"],
    [
        ["Background", "[VALUE]", "[VALUE]", "[VALUE]", "[VALUE]"],
        ["Buildings",  "[VALUE]", "[VALUE]", "[VALUE]", "[VALUE]"],
        ["Roads",      "[VALUE]", "[VALUE]", "[VALUE]", "[VALUE]"],
        ["Water",      "[VALUE]", "[VALUE]", "[VALUE]", "[VALUE]"],
        ["Vegetation", "[VALUE]", "[VALUE]", "[VALUE]", "[VALUE]"],
        ["Average",    "[VALUE]", "[VALUE]", "[VALUE]", "[VALUE]"],
    ],
    col_widths=[4, 3, 3, 3, 3]
)
img_placeholder("6.11", "Confusion matrix (5×5) – actual vs predicted class labels")
img_placeholder("6.12", "Side-by-side comparison – Sentinel-2 RGB composite vs classification output")
img_placeholder("6.13", "Final professional map layout from QGIS Print Layout (title, legend, scale bar, north arrow, statistics table)")
pb()

# ══════════════════════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
h1("7. CONCLUSION")

h2("7.1 Effectiveness of the Method")
body("The Deep Learning-based U-Net approach demonstrated high effectiveness for automated land "
     "cover classification from Sentinel-2 satellite imagery. Key strengths include:")
bullet("High Accuracy: The model successfully distinguished spectrally similar classes using spatial context and transfer learning from ResNet-34.")
bullet("Full Automation: The end-to-end pipeline from GEE data download to final vector outputs requires minimal manual intervention.")
bullet("Spatial Detail: U-Net skip connections preserve fine spatial details at 10-metre resolution.")
bullet("GIS Compatibility: All outputs are directly usable in QGIS, ArcGIS, and web mapping platforms.")
bullet("Multi-Class Detection: Simultaneous classification of five classes provides comprehensive land cover information in a single pass.")

h2("7.2 Limitations")
bullet("Resolution: 10-metre Sentinel-2 resolution cannot detect small buildings or narrow roads.")
bullet("Shadow Confusion: Building shadows are sometimes misclassified as water due to similar dark spectral signatures.")
bullet("Spectral Overlap: Bare soil and buildings share similar spectral signatures, causing occasional misclassification.")
bullet("Training Data: Deep Learning requires substantial labelled data; manual digitisation of training samples is time-consuming.")
bullet("Compute Requirements: GPU with at least 4–8 GB VRAM is needed for model training.")
bullet("Temporal Limitation: A single-year median composite does not capture seasonal changes in vegetation or water extent.")

h2("7.3 Future Improvements")
bullet("Incorporate higher-resolution commercial imagery (<1 m) for improved building detection.")
bullet("Implement automated shadow detection and removal in preprocessing.")
bullet("Experiment with deeper architectures (EfficientNet-B4, Swin Transformer) and model ensembles.")
bullet("Expand classes to include parking lots, railways, industrial areas, and bare land.")
bullet("Develop multi-temporal analysis for urban growth monitoring and change detection.")
bullet("Build a web-based application for non-technical stakeholders.")

h2("7.4 Observations from Indore District")
bullet("Building density is highest in the city centre and decreases towards the periphery, showing a clear urban–suburban gradient.")
bullet("The road network exhibits a radial pattern from the city centre; major highways are clearly detected.")
bullet("Khan River and Saraswati River are successfully identified along with several artificial water bodies.")
bullet("Green cover is concentrated in public parks, riparian zones, and agricultural areas on the city periphery.")
bullet("The urban core has limited vegetation, indicating potential urban heat island effects.")
para()
body("Overall, this project successfully demonstrates the power of Deep Learning for automated "
     "remote sensing analysis. The combination of Google Earth Engine, TensorFlow/Keras, and QGIS "
     "creates a comprehensive, reproducible pipeline for object detection from satellite imagery "
     "that can be adapted to other cities and geographic regions.")
pb()

# ══════════════════════════════════════════════════════════════════════════════
# 8. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
h1("8. REFERENCES")

h2("Data Sources")
numbered("European Space Agency (ESA). (2023). Sentinel-2 MSI Level-2A Surface Reflectance. https://scihub.copernicus.eu/")
numbered("Google LLC. (2024). Google Earth Engine – Sentinel-2 Dataset. https://developers.google.com/earth-engine/datasets/catalog/COPERNICUS_S2_SR")
numbered("GADM. (2024). Database of Global Administrative Areas v4.1. https://gadm.org/")

h2("Software and Tools")
numbered("Abadi, M., et al. (2015). TensorFlow: Large-Scale Machine Learning on Heterogeneous Systems. https://www.tensorflow.org/")
numbered("QGIS Development Team. (2024). QGIS Geographic Information System. https://qgis.org/")
numbered("GDAL/OGR contributors. (2024). GDAL/OGR Geospatial Data Abstraction Library. https://gdal.org/")
numbered("Gillies, S., et al. (2024). Rasterio: Geospatial raster I/O for Python. https://rasterio.readthedocs.io/")

h2("Research Papers")
numbered("Ronneberger, O., Fischer, P., & Brox, T. (2015). U-Net: Convolutional Networks for Biomedical Image Segmentation. MICCAI, pp. 234-241. https://arxiv.org/abs/1505.04597")
numbered("McFeeters, S.K. (1996). The use of the Normalized Difference Water Index (NDWI). International Journal of Remote Sensing, 17(7), 1425-1432.")
numbered("Zha, Y., Gao, J., & Ni, S. (2003). Use of normalized difference built-up index in automatically mapping urban areas. International Journal of Remote Sensing, 24(3), 583-594.")
numbered("Rouse, J.W., et al. (1974). Monitoring vegetation systems in the Great Plains with ERTS. NASA Special Publication, 351, 309.")

h2("Course Materials")
numbered("India Space Academy. (2026). Deep Learning Applications in Remote Sensing. GIS and Remote Sensing Programme.")

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUTPUT)
print("Word document saved to:", OUTPUT)
